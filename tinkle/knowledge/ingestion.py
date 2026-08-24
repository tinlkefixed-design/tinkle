from __future__ import annotations
from pathlib import Path
from typing import Any

from tinkle.knowledge.schemas import KnowledgeDocument

class DocumentIngestor:
    """Parse supported document types into a normalized KnowledgeDocument.

    Optional parsers are imported lazily so the core can still run without every
    document dependency installed. Unsupported formats fail explicitly.
    """
    SUPPORTED = {".txt", ".md", ".csv", ".pdf", ".docx", ".xlsx", ".xlsm"}

    def ingest_path(self, path: str | Path, *, metadata: dict[str, Any] | None = None) -> KnowledgeDocument:
        p = Path(path).expanduser().resolve()
        if not p.is_file():
            raise FileNotFoundError(str(p))
        suffix = p.suffix.lower()
        if suffix not in self.SUPPORTED:
            raise ValueError(f"Unsupported document type: {suffix or '<none>'}")
        content, parser_metadata = self._parse(p, suffix)
        merged = {"filename": p.name, "extension": suffix, **parser_metadata, **(metadata or {})}
        return KnowledgeDocument(title=p.stem, source=str(p), content=content, metadata=merged)

    def _parse(self, path: Path, suffix: str) -> tuple[str, dict[str, Any]]:
        if suffix in {".txt", ".md", ".csv"}:
            return path.read_text(encoding="utf-8", errors="replace"), {"parser": "text"}
        if suffix == ".pdf":
            import fitz
            parts: list[str] = []
            with fitz.open(path) as pdf:
                for number, page in enumerate(pdf, start=1):
                    text = page.get_text("text").strip()
                    if text:
                        parts.append(f"[Page {number}]\n{text}")
            return "\n\n".join(parts), {"parser": "pymupdf", "pages": len(parts)}
        if suffix == ".docx":
            from docx import Document
            doc = Document(path)
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
                if rows:
                    parts.append("\n".join(rows))
            return "\n\n".join(parts), {"parser": "python-docx", "paragraphs": len(doc.paragraphs)}
        if suffix in {".xlsx", ".xlsm"}:
            from openpyxl import load_workbook
            wb = load_workbook(path, read_only=True, data_only=True)
            sheets: list[str] = []
            parts: list[str] = []
            try:
                for ws in wb.worksheets:
                    sheets.append(ws.title)
                    rows = []
                    for row in ws.iter_rows(values_only=True):
                        vals = ["" if v is None else str(v) for v in row]
                        if any(vals):
                            rows.append(" | ".join(vals))
                    if rows:
                        parts.append(f"[Sheet: {ws.title}]\n" + "\n".join(rows))
            finally:
                wb.close()
            return "\n\n".join(parts), {"parser": "openpyxl", "sheets": sheets}
        raise ValueError(f"Unsupported document type: {suffix}")
